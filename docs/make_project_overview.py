"""Generate the plain-English project overview as .docx (python-docx).

Usage: .venv/bin/python docs/make_project_overview.py
Output: docs/Genesis_Engine_Project_Overview_2026-06-12.docx
Source facts: on-disk run reports + campaign RESULTS files (digested 2026-06-12).
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).with_name("Genesis_Engine_Project_Overview_2026-06-12.docx")

doc = Document()

# ---- base styling ----------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def title(text, sub=None, date=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(26)
    r.font.bold = True
    if sub:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(sub)
        r2.font.size = Pt(13)
        r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    if date:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(date)
        r3.font.size = Pt(11)
        r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def para(text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)


def gloss(term, definition):
    p = doc.add_paragraph()
    r = p.add_run(term + " — ")
    r.bold = True
    p.add_run(definition)


# ---- title page ------------------------------------------------------------
title(
    "The Genesis Engine Project",
    sub="Teaching a computer to invent board games — a plain-English overview",
    date="12 June 2026",
)
doc.add_paragraph()

# ---- at a glance ----------------------------------------------------------
h1("At a Glance")
table = doc.add_table(rows=0, cols=2)
table.style = "Light Grid Accent 1"
for k, v in [
    ("The goal", "Get a computer to invent a genuinely new board game that AI "
     "judging panels rate above 5 out of 10 (“I’d happily play this "
     "again”). No machine-invented game has cleared that bar yet."),
    ("Best honest score so far", "4.375 / 10 campaign average (Run 19); the "
     "long-standing benchmark game sits at 4.10 / 10."),
    ("Where we are", "The original “autopilot” approach hit a ceiling and "
     "was retired on pre-agreed terms. We now test rule ideas one at a time, "
     "clinical-trial style. The next candidate — “Frontline” — is "
     "fully built, has passed all pre-flight checks, and is ready to run."),
    ("One open decision", "A quirk of the current board lets a copy-cat "
     "(“mirror”) strategy avoid losing. We either trust trained AIs to "
     "punish it, or switch to a board size that makes mirroring impossible."),
]:
    row = table.add_row()
    row.cells[0].paragraphs[0].add_run(k).bold = True
    row.cells[1].text = v
doc.add_paragraph()

# ---- 1. the goal ------------------------------------------------------------
h1("1. The Goal")
para(
    "Most new board games are invented by people, and most are variations on "
    "ideas we already know — a twist on chess, a cousin of Go. This project "
    "asks a simple question with a hard answer: can a computer invent a "
    "genuinely good, genuinely new board game on its own?"
)
para(
    "“Good” here has a specific meaning, borrowed from the classics: "
    "simple rules that produce deep strategy. Go has a handful of rules and "
    "centuries of unexhausted depth. We want the computer to rediscover that "
    "property — and ideally find a game no human would have designed."
)
para("The machine works like a breeding program, in four steps:")
bullet("Thousands of candidate games are “bred” — the computer "
       "mixes and mutates rules (board shapes, how pieces capture, how you win) "
       "to produce new game designs.", bold_lead="Generate. ")
bullet("Obviously broken games are thrown out — games where one side "
       "always wins, or where nothing ever happens.", bold_lead="Sanity-check. ")
bullet("For each surviving game, an AI teaches itself to play it from scratch "
       "by playing against itself thousands of times — so we judge games as "
       "played with some skill, not by button-mashing.", bold_lead="Train. ")
bullet("The most promising games go to independent panels of AI judges who "
       "actually play them, write strategy guides, and score them out of 10.",
       bold_lead="Judge. ")
para(
    "The finish line is a score above 5.0 from those judging panels — a bar "
    "that, after we corrected for early grade inflation, no machine-invented "
    "game has yet reached."
)

# ---- 2. achieved ------------------------------------------------------------
h1("2. What Has Been Achieved")
para(
    "Fifteen major breeding runs and eight focused experiments later, the "
    "project has produced four things of lasting value."
)
h2("A working game factory")
para(
    "The pipeline reliably invents, trains and play-tests roughly 500 new games "
    "per run, end to end, on a single desktop computer. Every game, training "
    "session and judging verdict is recorded and reproducible."
)
h2("Genuine discoveries")
bullet("A capture rule that produces “ko fights” — tense, repeating "
       "capture-recapture battles — through a mechanism that exists in no "
       "traditional game (Run 10). This remains the project’s clearest "
       "piece of genuine machine novelty.")
bullet("Proof that the shape of the board is a quality lever in its own right: "
       "boards built like fractals (sponge-like shapes with holes at every "
       "scale) measurably outperformed ordinary grids (Runs 18–19).")
bullet("A new way to win — controlling zones of influence rather than "
       "occupying squares — that blind judging panels consistently preferred "
       "over the previous generation of games (June 2026). The specific first "
       "version failed its formal bar, but the underlying idea was validated "
       "by every panel that saw it without knowing what it was.")
h2("An honest measurement system")
para(
    "Possibly the most important achievement is less glamorous: the project "
    "learned to stop fooling itself. Judging is now done blind (panels don’t "
    "know which game is the new candidate and which is the old benchmark), "
    "experiments are pre-registered (success criteria are written down and "
    "locked before any results exist, as in clinical trials), and every score "
    "is compared against a single fixed benchmark game so that grade inflation "
    "can’t creep back in. Failures are reported as faithfully as wins."
)
h2("AI judges as bug-hunters")
para(
    "An unexpected bonus: the judging panels turned out to be exceptional "
    "bug-finders. By actually playing the games, they uncovered more than nine "
    "deep flaws in the game engine that the automated statistics had missed "
    "entirely — including one bug that was quietly manufacturing a "
    "promising-looking trend that evaporated once fixed."
)

# ---- 3. where we are --------------------------------------------------------
h1("3. Where We Are Now")
para(
    "The project is at a deliberate turning point, reached by rule rather than "
    "by mood."
)
para(
    "For its first phase, the breeding program was steered by an automated "
    "quality score called Go Essence — a formula meant to estimate "
    "“strategic depth per unit of rule complexity.” It worked well "
    "enough to get started, but over time a pattern emerged: the games the "
    "formula loved were not the games the judging panels loved. By Run 21 the "
    "two had become actively opposed at the extremes — the formula’s "
    "favourite game ranked near the bottom with judges, and a game the formula "
    "scored at almost zero tied for first. Scores had plateaued at 3.5–3.8 "
    "out of 10 for four consecutive runs.",
    bold_lead="The plateau. ",
)
para(
    "Crucially, the project had pre-committed to a stopping rule: two "
    "consecutive sub-5.0 campaigns would mean abandoning the autopilot. The "
    "rule fired in June 2026, and we honoured it.",
)
para(
    "Since the pivot, rule ideas are tested directly, one at a time, in a "
    "fixed sequence: cheap arithmetic checks first, then a calibration step, "
    "then a mechanical screen, then blind judging — with the criteria for "
    "passing each stage locked in advance. Two ideas have been through the "
    "full process and formally failed their bars (an influence-based win "
    "condition in its first parameterization, and a game with asymmetric "
    "goals where one side hunts the other). Both failures were cheap, "
    "informative, and taught us exactly which ingredients work: the influence "
    "idea itself, and a “flip” capture rule that converts enemy pieces "
    "rather than removing them.",
    bold_lead="The new method. ",
)
para(
    "The third candidate — Frontline — combines everything learned so "
    "far. Its core idea: you only score points for territory that both players "
    "are actively contesting. Quietly building in your own corner is worth "
    "exactly zero, so the game forces the head-to-head interaction that earlier "
    "machine-invented games lacked. The complete test apparatus is built, "
    "reviewed, and has passed all of its pre-flight checks. Notably, the "
    "design’s biggest historical flaw — captures that hurt the "
    "capturer — was proven structurally impossible in the final rules.",
    bold_lead="Frontline, ready to run. ",
)
para(
    "One decision is pending before the campaign starts. The pre-flight checks "
    "confirmed that on the current 22×22 board, a player who simply mirrors "
    "every opponent move through the centre can guarantee at least a draw. We "
    "must choose: keep the board and rely on trained AIs to punish mirroring "
    "(the formal test for this is already built in), or switch to a 21×21 "
    "board, whose centre square makes perfect mirroring impossible — at the "
    "cost of half a day’s rework and slightly weaker comparability with "
    "the benchmark games.",
    bold_lead="The mirror decision. ",
)
para(
    "In parallel, a second research track is hunting for a replacement for the "
    "retired Go Essence score — an automated signal that genuinely tracks "
    "what judges value. The first candidate (“drama” — how often "
    "the eventual winner was behind) looked promising in early tests but "
    "failed honestly when optimized against: the system learned to produce "
    "artificially close games rather than good ones. That failure was caught "
    "by the same pre-registered discipline, and better candidates are queued.",
    bold_lead="The selection-signal problem. ",
)

# ---- 4. hurdles --------------------------------------------------------------
h1("4. Hurdles, and How We Tackled Them")
h2("Grade inflation")
para(
    "Early in the project a champion game scored 8/10, and that number anchored "
    "expectations for months. When it was re-judged a few months later under "
    "the matured, consistent rubric, it scored 4.10. Nothing about the game had "
    "changed — the yardstick had. Fix: one benchmark game is now re-judged "
    "alongside every experiment, and all results are read as distances from "
    "that anchor rather than as raw numbers."
)
h2("Teaching to the test (Goodhart’s law)")
para(
    "Any time you optimize hard against a proxy measure, you get things that "
    "score well on the proxy rather than things that are actually good. The "
    "breeding program did exactly this to the Go Essence formula — it "
    "produced “packing race” games where both players quietly fill "
    "space without interacting, which the formula liked and judges hated. The "
    "same disease later appeared with the “drama” signal. Fix: "
    "automated scores are now diagnostic only; nothing gets selected on a "
    "number that hasn’t been validated against blind judge verdicts, and "
    "anything that will be optimized against gets adversarially stress-tested "
    "first."
)
h2("Luck masquerading as skill")
para(
    "AI training has a random element: the same game trained twice can look "
    "brilliant once and broken once. Early runs were systematically flattered "
    "by lucky training runs — in Run 21, a “phantom champion” was "
    "caught live when re-verification collapsed its score from 0.417 to 0.095. "
    "Fix: every score is now an average over multiple independent training "
    "runs with fixed random seeds, and champions are re-verified before being "
    "believed."
)
h2("Bugs that manufacture discoveries")
para(
    "Several times, an exciting trend turned out to be a software bug. The "
    "most instructive: a subtle bias in how simultaneous moves were processed "
    "made one whole category of games look promising; seven independent "
    "judging panels each found the bug by playing. Fix: a regression suite of "
    "400+ automated tests, a hard guarantee that new features cannot change "
    "the behaviour of existing games (verified bit-for-bit), and judges kept "
    "in the loop precisely because they catch what statistics miss."
)
h2("Wishful analysis")
para(
    "The subtlest hurdle was the temptation to move the goalposts after seeing "
    "results — to re-parameterize one more time, to explain away a bad "
    "number. Fix: pre-registration. Every experiment’s success bars, "
    "failure bars and permitted follow-ups are committed to the repository "
    "before any data exists, and the project does what the locked document "
    "says. This has now retired two of its own favourite ideas — painful "
    "but correct. In the same spirit, every new design is attacked by "
    "independent adversarial reviewers before money is spent on it; that "
    "process caught two fatal design flaws in Frontline (a do-nothing strategy "
    "that couldn’t lose, and the mirror exploit) while they were still "
    "cheap to fix on paper."
)
h2("Games that refuse to interact")
para(
    "The deepest design problem, finally named by the Run 21 judging panels: "
    "most machine-bred games let players succeed without ever engaging each "
    "other — parallel solitaire. Everything since has attacked this "
    "directly: win conditions built on influence and connection (validated "
    "blind), capture rules that convert rather than remove (validated blind, "
    "+0.2 over benchmark), and now Frontline, where interaction is literally "
    "the only source of points."
)

# ---- 5. next steps -----------------------------------------------------------
h1("5. Next Steps")
bullet("Make the mirror decision (keep the 22×22 board, or switch to "
       "21×21). Half a day either way; the runbook covers both branches.",
       bold_lead="1. ")
bullet("Run the Frontline campaign: calibration (~3 hours of computer time), "
       "mechanical screen (~2 hours), then blind judging by two independent "
       "panels. Success is pre-defined: beat the benchmark by a full point, "
       "which would be the first plateau break in five campaigns.",
       bold_lead="2. ")
bullet("If Frontline succeeds, it becomes the foundation game family for a "
       "new breeding era — this time selected by signals validated against "
       "judges, not by the retired formula. If it fails its bars, the "
       "rules-track closes honestly and effort concentrates on the "
       "selection-signal research.", bold_lead="3. ")
bullet("Either way, continue the search for a trustworthy automated quality "
       "signal (current candidates: how far ahead a player must plan; how "
       "quickly a learner improves), and then revive large-scale search with "
       "the archive machinery that has already passed its own validation.",
       bold_lead="4. ")

# ---- appendix A ---------------------------------------------------------------
doc.add_page_break()
h1("Appendix A — Run-by-Run Overview")
para(
    "Scores are out of 10 and, unless marked otherwise, come from independent "
    "AI judging panels. Scores before the May 2026 recalibration are inflated "
    "by roughly 4 points relative to today’s yardstick."
)

RUNS = [
    ("Run 7 (Feb 2026)",
     "First full pipeline test: breed games, train AIs, score, judge. The "
     "champion (“3D Forced-Capture Go”) scored 8/10 on the early, "
     "generous rubric and had a genuinely novel mechanic — but only 1 of "
     "the top 10 games had real merit, and a scoring bug read “non-"
     "triviality” as zero everywhere. Lesson: the factory works; the "
     "quality score doesn’t."),
    ("Run 8 (Mar 2026)",
     "Infrastructure fixes (better parameter bounds, fair seat-swapping in "
     "evaluation). Quality scores jumped 17-fold and the champion, "
     "“Connection Go,” scored 8/10 — the project’s "
     "long-standing benchmark game. Re-judged in May 2026 under the consistent "
     "rubric: 4.10. Lesson: boring fixes beat clever algorithms."),
    ("Run 9 (Mar 2026)",
     "Bigger population, same rules. Champion 7/10, but the entire top-20 were "
     "minor variations of “Go on a hexagonal board” — the search "
     "had found a comfortable rut. Lesson: more searching doesn’t help if "
     "the vocabulary of rules is too narrow."),
    ("Run 10 (Mar 2026)",
     "Added new board shapes and piece movement. Three games tied at 7/10, and "
     "one produced “ko fights” — repeating capture battles — "
     "via a mechanism found in no traditional game: the project’s clearest "
     "genuine discovery. Lesson: modest new building blocks unlock real "
     "novelty."),
    ("Run 11 (Mar 2026)",
     "A consolidation run: the failure modes discovered in Run 10 (e.g. a "
     "board shape that made one win condition trivial) were folded into the "
     "automatic sanity checks. No standalone judging round."),
    ("Run 12 (Mar 2026)",
     "First serious attempt at “cellular automata” games — rules "
     "where the board itself evolves like Conway’s Game of Life. Outcome: "
     "zero CA games in the top 20, but the autopsy found three specific, "
     "fixable causes. Lesson: a useful failure is worth more than a vague "
     "success."),
    ("Run 13 (Apr 2026)",
     "CA fixes applied; first large judging campaign (22 verdicts) with a "
     "dedicated “novelty adversary.” CA games became competitive, but "
     "the formula’s favourite scored 5.0 while a game it ranked third "
     "won at 5.9 — the first clear sign the formula and the judges "
     "disagree. A judging panel also found a deep, year-old distance bug. "
     "Lesson: judges overrule metrics, and they find bugs."),
    ("Run 14 (Apr 2026)",
     "Added simultaneous play (both players move at once). No game beat its "
     "predecessors, and seven independent panels each discovered the same "
     "engine bias that had been fabricating the apparently promising "
     "simultaneous-play signal. Lesson: never trust a trend while a bias is "
     "live."),
    ("Run 15 (Apr 2026)",
     "Five engine fixes shipped. The formula’s champion scored 2.43 — "
     "the worst ever — with a complete inversion between formula rank and "
     "judge rank. Lesson: runs 13–15’s real product was nine deep "
     "bug-fixes, none of which the formula could have found."),
    ("Run 16 (Apr 2026)",
     "More principled rule resolution and generator filters. Judge winner at "
     "4.40 — the best non-benchmark result of the era — and the "
     "simultaneous-play idea was conclusively retired. Lesson: the strongest "
     "rule family combines alternating turns, active capture, and influence."),
    ("Run 17 (Apr 2026)",
     "A clean-engineering run (audits, hard rules against imbalanced games, "
     "first fractal-board seed). Champion 4.14, run average 3.50. Lesson: "
     "clean engineering alone doesn’t buy depth — the honest signal "
     "said the games were mediocre."),
    ("Run 18 (Apr–May 2026)",
     "A systematic scan of board “dimension”: six board families from "
     "sparse fractals to the dense “menger sponge.” Higher-dimension "
     "fractal boards clearly outperformed — a real discovery — but "
     "much of the apparent generation-on-generation progress turned out to be "
     "lucky training seeds. Lesson: board shape matters; so does honest "
     "averaging."),
    ("Run 19 (May 2026)",
     "First run under the fully honest measurement stack (fixed seeds, "
     "multi-seed averages). Judging campaign average 4.375 — the best of "
     "the recalibrated era and still the number to beat. Lesson: the honest "
     "pipeline works."),
    ("Run 20 (May 2026)",
     "Tested whether the benchmark’s rule family (connection wins) was the "
     "secret of its quality, by seeding 12 connection games. Every one of them "
     "evolved away into “territory race” games within five "
     "generations; meanwhile the run produced the deepest single game ever "
     "measured (depth 0.894). Campaign average 3.73. Lesson: depth exists in "
     "the search space, but selection by formula doesn’t keep it."),
    ("R8 replay & Run 20.5 (May 2026)",
     "The old 8/10 benchmark was re-judged under the modern rubric: 4.10 ± "
     "1.14. All historical comparisons were recalibrated, and the fair "
     "seat-swap evaluation method used by all later experiments was "
     "established."),
    ("Run 21 (May–Jun 2026)",
     "The focusing run: deduplication, komi (a points handicap for the second "
     "player), live re-verification of champions — which caught a phantom "
     "champion whose score collapsed from 0.417 to 0.095 on re-test. Judging: "
     "no game cleared 5.0; campaign average 3.69; and the formula’s "
     "rankings were now inverted against the judges’ at both extremes. "
     "The pre-committed stopping rule fired: pivot. Lesson: stop evolving "
     "against the formula; test rule ideas directly."),
    ("Field-Connect probe (9 Jun 2026)",
     "First pre-registered experiment: a win condition based on zones of "
     "influence, judged blind against the plateau baseline. Formal result: "
     "failed its bar (+0.70 vs the required +1.0) — but both blind panels "
     "preferred it (4.15 vs 3.45) and independently praised the same "
     "ingredient: a win condition that forces interaction. Lesson: right "
     "lever, wrong tuning."),
    ("Phase 1.5 rules rethink (10 Jun 2026)",
     "Tested the panels’ top tuning suggestions (a sharper, shorter-range "
     "field) plus three capture variants. The suggestions made the game "
     "measurably worse — the experiment refuted its own premise and "
     "validated the original tuning. One capture variant (“flip” "
     "capture, converting enemy pieces) showed real life. Lesson: locked "
     "criteria protect you from plausible-sounding advice."),
    ("Pivot menu (10 Jun 2026)",
     "Five candidate directions were designed and attacked by independent "
     "adversarial review panels. Ranked verdict: run SIEGE (a hunter-vs-"
     "builder asymmetric game) with flip-capture as its control arm; keep "
     "Frontline as the registered next family; defer the two search-machinery "
     "candidates until their instruments are fixed."),
    ("SIEGE campaign (11 Jun 2026)",
     "The asymmetric game died early and cheaply: the “hunter” side "
     "never learned to play well enough to pass the pre-set skill gate (0 of "
     "9 configurations). The control arm — flip-capture on the proven "
     "influence board — ran the full course and beat the benchmark by "
     "+0.2, unanimously but far below the +1.0 bar. Lesson: asymmetric goals "
     "retired; flip-capture is real but not a breakthrough alone; next family: "
     "Frontline."),
    ("RC2 selection research (11 Jun 2026)",
     "A parallel track searching for a replacement quality signal. A "
     "“drama” measure (how often the winner was behind) passed its "
     "anchor tests, and the archive-based search machinery passed validation "
     "— but when actually optimized against, drama produced artificially "
     "close, bad games (the worst game any campaign has fielded). The signal "
     "was honestly retired as a target. Lesson: closeness is not depth; the "
     "machinery is ready, the signal is still missing."),
    ("Frontline build (11–12 Jun 2026)",
     "The registered next family, rebuilt from its previously fatal design: "
     "score only in territory both players contest. Spec and success criteria "
     "locked before any code; the build ran through 13 reviewed tasks; all "
     "pre-flight arithmetic and simulation checks passed — including a "
     "proof that the old “captures hurt the capturer” defect is now "
     "structurally impossible. One finding awaits a decision: a mirror "
     "strategy can force a draw on the current even-sized board. The campaign "
     "itself has not yet run."),
]
for name, text in RUNS:
    h2(name)
    para(text)

# ---- appendix B ---------------------------------------------------------------
doc.add_page_break()
h1("Appendix B — Glossary")
para("Plain-English definitions of the terms used in this document and in the "
     "project’s reports.")

GLOSSARY = [
    ("Abstract strategy game",
     "A board game with no luck, no hidden information and usually no theme "
     "— just rules and thinking. Chess, Go and Hex are the classics."),
    ("Agent / AI agent",
     "A computer program that perceives a situation and chooses actions. Here: "
     "either a trained game-playing program, or an AI (Claude) acting as a "
     "judge, reviewer or experimenter."),
    ("Anchor",
     "A fixed reference game (here, Run 8’s “Connection Go”, "
     "scoring 4.10) that is re-judged alongside every experiment so all "
     "results can be compared on one yardstick."),
    ("Blind evaluation",
     "Judging games without knowing which is the new candidate and which is "
     "the benchmark — labels are scrambled and the key is sealed until "
     "all verdicts are filed. Prevents bias, like a blind taste test."),
    ("Calibration (Stage 1)",
     "A cheap tuning step before the real experiment: checking a game is fair "
     "between the two players and ends in sensible ways, and picking its "
     "settings from a pre-agreed menu."),
    ("Campaign",
     "One complete pre-registered experiment, from locked criteria through "
     "calibration, screening and blind judging to a formal verdict."),
    ("Capture",
     "Any rule for removing or converting enemy pieces. “Surround” "
     "capture removes pieces that run out of breathing room (as in Go); "
     "“flip” capture converts a dominated enemy piece to your own "
     "colour (as in Reversi, but driven by influence here)."),
    ("Cellular automaton (CA)",
     "A rule that updates every cell of a board simultaneously based on its "
     "neighbours — like Conway’s Game of Life. Tried as a game "
     "ingredient in Runs 12–16."),
    ("Comparator / control arm",
     "A known game run through an experiment alongside the new candidate, so "
     "differences can be attributed to the one thing that changed."),
    ("Contested-majority scoring",
     "Frontline’s core rule: a cell only scores if both players project "
     "meaningful influence onto it — you get points only for territory "
     "you are actually fighting over."),
    ("Decision rule / grammar (GO, PARTIAL, NO-GO)",
     "The pre-agreed mapping from results to actions: GO = the idea passed its "
     "bar, build on it; PARTIAL = one specific licensed retry; NO-GO = retire "
     "the idea. Written down before the experiment runs."),
    ("Drama (signal)",
     "A candidate quality measure: how far behind the eventual winner was "
     "during the game. Promising as a diagnostic; retired as a target after "
     "it was shown to reward artificially close games."),
    ("Engine",
     "The software that knows the rules: it hosts the board, enforces legal "
     "moves, applies captures and decides who has won."),
    ("Evolutionary search",
     "Improving designs the way breeding improves crops: keep the best "
     "performers, produce variations of them (mutation) and combinations of "
     "them (crossover), repeat for many generations."),
    ("Fitness function",
     "The score an evolutionary search tries to maximize. For most of this "
     "project that was Go Essence; choosing a fitness function that matches "
     "what you actually want is the hard part."),
    ("Fractal board",
     "A game board with holes at every scale (like a sponge or a snowflake), "
     "giving it an effective dimension between the usual 2 and 3. Higher-"
     "dimension fractal boards measurably improved game quality."),
    ("Go Essence (GE)",
     "The project’s original automated quality formula: roughly, "
     "strategic depth times variety, divided by rule complexity. Useful "
     "early; retired as a selection signal in June 2026 after its rankings "
     "inverted against judge verdicts."),
    ("Goodhart’s law",
     "“When a measure becomes a target, it ceases to be a good "
     "measure.” The recurring villain of this project: optimize hard "
     "against any proxy score and you get things that game the score."),
    ("Hex",
     "A classic connection game: players race to link their two opposite "
     "board edges. Several of the project’s games are descendants of "
     "this idea."),
    ("Influence field",
     "A map of board control: every stone radiates influence that fades with "
     "distance (full strength on its own cell, half next door, a quarter two "
     "cells away). Cells are controlled by whoever projects more influence."),
    ("KILL bar",
     "A pre-registered tripwire that stops an experiment early and cheaply if "
     "a design is fundamentally broken — e.g. “if captures never "
     "fire at all, stop here.”"),
    ("Komi",
     "A small points bonus given to the second player to offset the advantage "
     "of moving first — borrowed from Go."),
    ("MAP-Elites / Quality-Diversity (QD)",
     "A modern search method that keeps an archive of the best design in each "
     "“niche” (e.g. best short game, best capture-heavy game) instead "
     "of a single champion — so the search explores broadly rather than "
     "converging on one rut. The machinery is validated and waiting for a "
     "trustworthy quality signal."),
    ("Mirror strategy",
     "Copying the opponent’s every move reflected through the board’s "
     "centre. On an even-sized board with no centre cell this can guarantee "
     "symmetry — and at least a draw — which is why a 21×21 "
     "switch is on the table for Frontline."),
    ("Mutation / crossover",
     "The two ways evolution makes new designs: small random changes to one "
     "design (mutation) and mixing parts of two designs (crossover)."),
    ("Pie rule",
     "A fairness device: player one makes the first move, then player two may "
     "either reply normally or swap sides and take that move as their own. "
     "Discourages overpowering first moves — like one person cutting the "
     "cake and the other choosing the slice."),
    ("Plateau",
     "The project’s four-run stretch (Runs 18–21 era) where judged "
     "quality stalled at 3.5–3.8 out of 10, triggering the pre-agreed "
     "pivot."),
    ("PPO (Proximal Policy Optimization)",
     "The standard reinforcement-learning algorithm used to train the game-"
     "playing AIs. Details aside: it learns by playing, gradually preferring "
     "actions that lead to wins."),
    ("Pre-registration",
     "Writing down — and locking in version control — an "
     "experiment’s success bars, failure bars and permitted follow-ups "
     "before any results exist. Borrowed from clinical trials; the project’s "
     "main defence against fooling itself."),
    ("Reinforcement learning (RL)",
     "Teaching by trial and error: the program tries actions, receives rewards "
     "(here: winning), and adjusts to earn more reward."),
    ("Run",
     "One complete cycle of the breeding program: generate hundreds of games, "
     "train and score them across several generations, then send the best to "
     "judging panels. The project has completed Runs 7–21."),
    ("Seat balance / first-mover advantage",
     "Whether going first confers an unfair edge. Measured for every game; "
     "fixed with the pie rule and komi where needed."),
    ("Self-play",
     "Training method where the AI improves by playing against copies of "
     "itself — no human examples needed."),
    ("Smoke test",
     "A quick, cheap check that something basic isn’t broken before "
     "spending real time and compute — named after switching a device on "
     "to see if smoke comes out."),
    ("Stages 0–3",
     "The fixed ladder every campaign climbs: Stage 0 — arithmetic and "
     "simulation sanity checks; Stage 1 — calibration; Stage 2 — "
     "mechanical screen against pre-set measurements; Stage 3 — blind "
     "judging by independent panels. Each stage can stop the campaign "
     "cheaply."),
    ("Threshold race / packing race",
     "The degenerate style of play the old formula accidentally rewarded: "
     "both players quietly accumulate territory in parallel without "
     "interacting — efficient, and dull."),
    ("Training budget",
     "How long each game’s AI players are trained (counted in learning "
     "steps). Calibration uses a smaller budget; final screens use a larger "
     "one."),
    ("Win condition",
     "The rule that says how a game ends and who wins — connect two "
     "sides, control the most territory, hold a lead long enough, etc. The "
     "project’s central finding is that this choice, more than any "
     "other, determines whether a game forces interesting interaction."),
]
for term, definition in sorted(GLOSSARY, key=lambda t: t[0].lower()):
    gloss(term, definition)

doc.save(OUT)
print(f"wrote {OUT} ({OUT.stat().st_size:,} bytes)")
